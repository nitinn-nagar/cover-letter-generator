import streamlit as st
from docx import Document
from io import BytesIO
import requests
import time

# Load CloudConvert API key securely
CLOUDCONVERT_API_KEY = st.secrets["cloudconvert"]["api_key"]

def process_template(doc, client_info):
    """Replaces placeholders with client info throughout the docx template, including headers/footers."""
    replacements = {
        "<<CLIENT_NAME>>": client_info['name'],
        "<<COMPANY>>": client_info['company'],
        "<<ADDRESS>>": f"{client_info['address1']}\n{client_info['address2']}" if client_info['address2'] else client_info['address1'],
        "<<ADDRESS_LINE_1>>": client_info['address1'],
        "<<ADDRESS_LINE_2>>": client_info['address2'],
        "<<DATE>>": client_info['date']
    }

    def replace_text(paragraph):
        for run in paragraph.runs:
            for key, val in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, val)

    def process_paragraphs(paragraphs):
        for para in paragraphs:
            replace_text(para)

    def process_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

    # Replace in body
    process_paragraphs(doc.paragraphs)
    process_tables(doc.tables)

    # Replace in headers & footers
    for section in doc.sections:
        # Header
        process_paragraphs(section.header.paragraphs)
        process_tables(section.header.tables)

        # Footer
        process_paragraphs(section.footer.paragraphs)
        process_tables(section.footer.tables)

    return doc


def convert_docx_to_pdf_cloudconvert(docx_bytes):
    """Uses CloudConvert API to convert DOCX to PDF with formatting preserved."""
    headers = {
        "Authorization": f"Bearer {CLOUDCONVERT_API_KEY}",
        "Content-Type": "application/json"
    }

    # Step 1: Create job
    job_resp = requests.post(
        "https://api.cloudconvert.com/v2/jobs",
        headers=headers,
        json={
            "tasks": {
                "upload-file": {
                    "operation": "import/upload"
                },
                "convert-file": {
                    "operation": "convert",
                    "input": "upload-file",
                    "input_format": "docx",
                    "output_format": "pdf",
                    "engine": "office"
                },
                "export-file": {
                    "operation": "export/url",
                    "input": "convert-file"
                }
            }
        }
    )

    job_data = job_resp.json()["data"]
    upload_url = job_data["tasks"][0]["result"]["form"]["url"]
    upload_params = job_data["tasks"][0]["result"]["form"]["parameters"]

    # Step 2: Upload DOCX file
    files = {"file": ("input.docx", docx_bytes)}
    requests.post(upload_url, data=upload_params, files=files)

    # Step 3: Poll job status
    job_id = job_data["id"]
    status = "processing"

    while status not in ["finished", "error"]:
        job_status = requests.get(f"https://api.cloudconvert.com/v2/jobs/{job_id}", headers=headers).json()
        status = job_status["data"]["status"]
        time.sleep(1)

    if status == "error":
        raise Exception("Conversion job failed")

    # Step 4: Get export URL
    export_task = [t for t in job_status["data"]["tasks"] if t["name"] == "export-file"][0]
    file_url = export_task["result"]["files"][0]["url"]
    pdf_data = requests.get(file_url).content
    return pdf_data

def main():
    st.set_page_config(page_title="Cover Letter Generator", page_icon="📝")
    st.title("📝 Cover Letter Generator")

    st.header("1. Enter Client Information")
    col1, col2 = st.columns(2)

    with col1:
        name = st.text_input("Client Name")
        company = st.text_input("Company Name")
        date = st.date_input("Date")
    with col2:
        address1 = st.text_input("Address Line 1")
        address2 = st.text_input("Address Line 2 (optional)")

    client_info = {
        "name": name,
        "company": company,
        "address1": address1,
        "address2": address2,
        "date": date.strftime("%B %d, %Y")

    }

    st.header("2. Upload DOCX Template")
    template_file = st.file_uploader("Upload .docx file", type="docx")

    output_format = st.radio("3. Choose output format:", ["DOCX", "PDF"], horizontal=True)

    st.markdown("#### 📌 Placeholder Tags in Your Template")
    st.code("<<CLIENT_NAME>>, <<COMPANY>>, <<ADDRESS>>, <<DATE>> \n Please place the placeholders in respective places in the word template. \n Align <<COMPANY>> and <<ADDRESS>> to the left. ")

    if st.button(f"🚀 Generate {output_format}"):
        if not all([name, company, address1, template_file]):
            st.error("Please fill in all required fields and upload a template.")
            return

        try:
            doc = Document(template_file)
            result_doc = process_template(doc, client_info)

            if output_format == "DOCX":
                output = BytesIO()
                result_doc.save(output)
                output.seek(0)
                st.download_button("⬇️ Download DOCX", output,
                    file_name=f"{name}_{company}_cover_letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                buffer = BytesIO()
                result_doc.save(buffer)
                buffer.seek(0)
                pdf = convert_docx_to_pdf_cloudconvert(buffer.read())
                st.download_button("⬇️ Download PDF", pdf,
                    file_name=f"{name}_{company}_cover_letter.pdf",
                    mime="application/pdf")

            st.success(f"{output_format} generated successfully!")

        except Exception as e:
            st.error("An error occurred while generating the cover letter.")
            st.exception(e)

main()
